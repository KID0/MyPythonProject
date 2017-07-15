# -*- coding: UTF-8 -*-
# 作者：刘浩

'''
脚本作用：
1.在本文件夹中找到所有的docx文件
2.从docx文件中提取所有段落，每遇到一个标点符号就进行分行
3.将这些数据打印到新的docx中
'''

# 使用方法：建立一个新文件夹，将需要处理的docx文件和该脚本放进去，运行该脚本即可
# 帮助文档：https://python-docx.readthedocs.io/en/latest/
# 参考文献：http://blog.csdn.net/qianchenglenger/article/details/51582005




def transform(NameOfDocx):
	'''
	输入一个docx文件，遇到["!","?",".",",",'！','。','？','，']进行分行
	将分行后的内容保存为“原文件名+revised”，也是docx格式

	'''

	'''
	python-docx组织结构：
	document
	  --heading
	  --paragraphs(a list)
	    --paragraph1
	    --paragraph2
	    --paragraph3
	    --......
	  --picture
	  --table

	'''
	from docx import Document

	document_origin = Document(NameOfDocx)
	document_new = Document()
	for paragraph in document_origin.paragraphs:
		sentence = []
		for word in paragraph.text:
			# 注意：这里不能用and or ，这样是布尔运算
			if word in ["!","?",".",",",'！','。','？','，']:
				sentence.append(word)
				document_new.add_paragraph(''.join(sentence))
				sentence = []
			else:
				sentence.append(word)
	NewNameOfDocx = NameOfDocx.split('.')[0] + ' revised.docx'
	document_new.save(NewNameOfDocx)
	return


def FindDocx():
	'''
	找出所有的.docx文件，存放在 target_file 中,并返回target_file列表
	'''

	# os.getcwd() 当前工作目录
	# os.listdir() 列出目标目录下的所有文件
	import os
	filenames = os.listdir(os.getcwd())
	target_file = []
	for filename in filenames:
		if filename.split('.')[-1] == 'docx':
			target_file.append(filename)
	return target_file

# 以下为运行主程序
target_file = FindDocx()
for file in target_file:
	transform(file)